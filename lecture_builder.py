#!/usr/bin/env python3
"""高中物理講義自動排版工具。

功能：
1. 讀取多張圖片並進行 OCR。
2. 依章節整理成 Word（.docx）講義。
3. 提供基礎美編樣式。
4. 若有配置 Mathpix API，優先使用可保留 LaTeX 公式的辨識結果。
"""

from __future__ import annotations

import argparse
import base64
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

import requests
import pytesseract
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


@dataclass
class OcrResult:
    source: Path
    text: str
    mode: str


class OcrEngine:
    def __init__(self, lang: str = "chi_tra+eng") -> None:
        self.lang = lang
        self.mathpix_id = os.getenv("MATHPIX_APP_ID")
        self.mathpix_key = os.getenv("MATHPIX_APP_KEY")

    def run(self, image_path: Path) -> OcrResult:
        if self.mathpix_id and self.mathpix_key:
            text = self._run_mathpix(image_path)
            return OcrResult(source=image_path, text=text, mode="mathpix")

        text = self._run_tesseract(image_path)
        return OcrResult(source=image_path, text=text, mode="tesseract")

    def _run_tesseract(self, image_path: Path) -> str:
        with Image.open(image_path) as img:
            text = pytesseract.image_to_string(img, lang=self.lang)
        return normalize_text(text)

    def _run_mathpix(self, image_path: Path) -> str:
        encoded = base64.b64encode(image_path.read_bytes()).decode("utf-8")
        payload = {
            "src": f"data:image/{image_path.suffix.lstrip('.').lower()};base64,{encoded}",
            "formats": ["text", "data"],
            "data_options": {"include_latex": True},
            "ocr": ["math", "text"],
        }
        headers = {
            "app_id": self.mathpix_id,
            "app_key": self.mathpix_key,
            "Content-type": "application/json",
        }
        response = requests.post("https://api.mathpix.com/v3/text", json=payload, headers=headers, timeout=90)
        response.raise_for_status()
        data = response.json()
        text = data.get("text") or ""
        return normalize_text(text)


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def discover_images(paths: Iterable[str]) -> List[Path]:
    found: List[Path] = []
    for raw in paths:
        p = Path(raw)
        if p.is_file():
            found.append(p)
            continue

        expanded = list(Path().glob(raw))
        found.extend(x for x in expanded if x.is_file())

    unique = sorted(set(found))
    if not unique:
        raise FileNotFoundError("找不到可處理的圖片，請確認路徑或 glob pattern。")
    return unique


def setup_document(document: Document, title: str) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")
    normal.font.size = Pt(12)

    h1 = document.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")
    h1.font.size = Pt(18)

    title_para = document.add_paragraph(title)
    title_para.style = document.styles["Title"]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph("（由圖片自動轉寫與排版）")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()


def add_ocr_block(document: Document, result: OcrResult, index: int) -> None:
    document.add_heading(f"第 {index} 頁：{result.source.name}", level=1)

    note = document.add_paragraph()
    note.add_run("辨識引擎：").bold = True
    note.add_run(result.mode)

    for chunk in split_chunks(result.text):
        p = document.add_paragraph(chunk)
        p.paragraph_format.line_spacing = 1.5

    document.add_page_break()


def split_chunks(text: str) -> List[str]:
    chunks = [x.strip() for x in text.split("\n\n")]
    return [x for x in chunks if x]


def build_docx(image_paths: List[Path], output: Path, title: str, lang: str) -> None:
    document = Document()
    setup_document(document, title)
    engine = OcrEngine(lang=lang)

    for idx, path in enumerate(image_paths, start=1):
        result = engine.run(path)
        add_ocr_block(document, result, idx)

    document.save(output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="將高中物理講義圖片自動轉寫並輸出為 Word 檔。")
    parser.add_argument("images", nargs="+", help="圖片檔或 glob pattern，例如 'physics/*.jpg'")
    parser.add_argument("-o", "--output", default="physics_handout.docx", help="輸出的 Word 檔案路徑")
    parser.add_argument("--title", default="高中物理補習班講義", help="講義封面標題")
    parser.add_argument("--lang", default="chi_tra+eng", help="Tesseract 語言，例如 chi_tra+eng")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    images = discover_images(args.images)
    build_docx(images, Path(args.output), args.title, args.lang)
    print(f"完成！已輸出：{args.output}")


if __name__ == "__main__":
    main()
